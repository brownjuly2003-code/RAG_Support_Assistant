# ingestion/loader.py
"""
Загрузчик локальных документов.

Функции:
- рекурсивный обход директории;
- чтение файлов поддерживаемых форматов (TXT, MD, PDF, DOCX, HTML);
- вычисление метаданных: размер, дата изменения, хэш содержимого;
- возврат списка Document (совместимы с LangChain/LangGraph).

Отслеживание изменений:
- в metadata для каждого документа храним:
    - mtime (ISO строка),
    - size_bytes,
    - content_hash (sha256);
- дополнительно есть вспомогательный класс DocumentChangeTracker
  для сохранения/сравнения состояния между запусками.

ИСПОЛЬЗУЕТ ТОЛЬКО ЛОКАЛЬНЫЕ ФАЙЛЫ, БЕЗ ВНЕШНИХ API.
"""

from __future__ import annotations

import hashlib
import json
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

try:
    from langchain.schema import Document
except ImportError:
    # Минимальный fallback, если langchain не установлен.
    @dataclass
    class Document:  # type: ignore
        page_content: str
        metadata: Dict[str, Any]


try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}


class DocumentLoader:
    """Загружает документы из локальной папки."""

    def __init__(self, recursive: bool = True):
        """
        :param recursive: обходить ли поддиректории.
        """
        self.recursive = recursive

    # --- публичный метод ---

    def load_documents(
        self,
        folder_path: str | Path,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        Загружает все документы из указанной папки.

        Args:
            folder_path: путь к папке с документами (например, demo/docs).
            extensions: список расширений, которые нужно загружать.
                        По умолчанию — все SUPPORTED_EXTENSIONS.

        Returns:
            Список Document с полями:
            - page_content: текст документа;
            - metadata:
                - source: имя файла;
                - file_path: абсолютный путь;
                - mtime: время последней модификации (ISO);
                - size_bytes: размер файла;
                - content_hash: sha256 хэш содержимого;
                - format: расширение без точки.
        """
        base = Path(folder_path).expanduser().resolve()
        if not base.exists():
            raise FileNotFoundError(f"Папка не существует: {base}")
        if not base.is_dir():
            raise ValueError(f"Ожидалась папка, а не файл: {base}")

        if extensions is None:
            exts = SUPPORTED_EXTENSIONS
        else:
            exts = {("." + ext.lstrip(".")).lower() for ext in extensions}

        documents: List[Document] = []

        for file_path in self._iter_files(base, exts):
            try:
                text = self._read_file(file_path)
                if not text.strip():
                    continue
                meta = self._build_metadata(file_path, text)
                documents.append(Document(page_content=text, metadata=meta))
            except Exception as e:
                print(f"✗ Ошибка при чтении {file_path}: {e}")

        print(f"✓ Загружено документов: {len(documents)}")
        return documents

    # --- внутренние утилиты чтения файлов ---

    def _iter_files(self, base: Path, exts: set[str]) -> List[Path]:
        """Список файлов нужных расширений (рекурсивно или нет)."""
        results: List[Path] = []
        if self.recursive:
            for root, _, files in os.walk(base):
                for name in files:
                    path = Path(root) / name
                    if path.suffix.lower() in exts and not name.startswith("."):
                        results.append(path)
        else:
            for path in base.iterdir():
                if path.is_file() and path.suffix.lower() in exts and not path.name.startswith("."):
                    results.append(path)
        return results

    def _read_file(self, path: Path) -> str:
        """Определяет метод чтения по расширению и возвращает текст."""
        ext = path.suffix.lower()
        if ext in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == ".pdf":
            return self._read_pdf(path)
        if ext == ".docx":
            return self._read_docx(path)
        if ext in {".html", ".htm"}:
            return self._read_html(path)
        raise ValueError(f"Неподдерживаемый формат: {ext}")

    def _read_pdf(self, path: Path) -> str:
        if pypdf is None:
            raise ImportError("Для чтения PDF установите pypdf: pip install pypdf")
        text_parts: List[str] = []
        with path.open("rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _read_docx(self, path: Path) -> str:
        if DocxDocument is None:
            raise ImportError("Для чтения DOCX установите python-docx: pip install python-docx")
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text)

    def _read_html(self, path: Path) -> str:
        """Простейшее извлечение текста из HTML без сторонних парсеров."""
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self._parts: List[str] = []

            def handle_data(self, data: str) -> None:
                data = data.strip()
                if data:
                    self._parts.append(data)

            def get_text(self) -> str:
                return " ".join(self._parts)

        content = path.read_text(encoding="utf-8", errors="ignore")
        parser = TextExtractor()
        parser.feed(content)
        return parser.get_text()

    # --- метаданные и хэш ---

    @staticmethod
    def _compute_hash(text: str) -> str:
        """sha256 хэш содержимого (для отслеживания изменений)."""
        h = hashlib.sha256()
        h.update(text.encode("utf-8", errors="ignore"))
        return h.hexdigest()

    def _build_metadata(self, path: Path, text: str) -> Dict[str, Any]:
        stat = path.stat()
        return {
            "source": path.name,
            "file_path": str(path.resolve()),
            "mtime": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "size_bytes": stat.st_size,
            "content_hash": self._compute_hash(text),
            "format": path.suffix.lstrip(".").lower(),
        }


class DocumentChangeTracker:
    """
    Вспомогательный класс для сравнения состояний документов между запусками.

    Идея:
    - после загрузки документов сохраняем состояние в JSON;
    - при следующем запуске загружаем старое состояние и сравниваем.

    Смотрим на file_path + content_hash, чтобы понять, что изменилось.
    """

    @staticmethod
    def save_state(docs: List[Document], path: str | Path) -> None:
        state: Dict[str, Any] = {}
        for doc in docs:
            meta = doc.metadata
            key = meta.get("file_path") or meta.get("source")
            if not key:
                continue
            state[key] = {
                "mtime": meta.get("mtime"),
                "size_bytes": meta.get("size_bytes"),
                "content_hash": meta.get("content_hash"),
            }

        out = Path(path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    @staticmethod
    def load_state(path: str | Path) -> Dict[str, Any]:
        p = Path(path)
        if not p.exists():
            return {}
        return json.loads(p.read_text(encoding="utf-8"))

    @staticmethod
    def diff(
        docs: List[Document],
        previous_state: Dict[str, Any],
    ) -> Dict[str, List[str]]:
        """
        Возвращает:
        - new: новые файлы;
        - modified: изменившиеся (по content_hash);
        - deleted: отсутствующие сейчас, но бывшие раньше.
        """
        current: Dict[str, str] = {}
        for doc in docs:
            meta = doc.metadata
            key = meta.get("file_path") or meta.get("source")
            if key:
                current[key] = meta.get("content_hash", "")

        changes = {"new": [], "modified": [], "deleted": []}

        # новые и изменённые
        for key, chash in current.items():
            prev = previous_state.get(key)
            if prev is None:
                changes["new"].append(key)
            elif prev.get("content_hash") != chash:
                changes["modified"].append(key)

        # удалённые
        for key in previous_state.keys():
            if key not in current:
                changes["deleted"].append(key)

        return changes


if __name__ == "__main__":
    # Пример ручного запуска
    demo_dir = Path(__file__).parent.parent / "demo" / "docs"
    loader = DocumentLoader()
    docs = loader.load_documents(demo_dir)

    tracker = DocumentChangeTracker()
    state_file = Path(__file__).parent.parent / "data" / "loader_state.json"
    tracker.save_state(docs, state_file)
    prev = tracker.load_state(state_file)
    changes = tracker.diff(docs, prev)
    print("Изменения:", changes)
