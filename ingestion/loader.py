"""
ingestion/loader.py

DocumentLoader: loads documents from a directory or a single file.

Supported formats:
    .txt, .md, .pdf (pypdf), .docx (python-docx), .json, .csv, .html, .htm

Each document gets metadata:
    - source: filename
    - file_path: absolute path
    - file_type: extension without dot
    - page: page number (for multi-page PDFs, one Document per page)
    - mtime: last modified ISO timestamp
    - size_bytes: file size
    - content_hash: sha256 of the extracted text
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from langchain_core.documents import Document
except ImportError:
    try:
        from langchain.schema import Document  # type: ignore[no-redef]
    except ImportError:
        from dataclasses import dataclass as _dc

        @_dc
        class Document:  # type: ignore[no-redef]
            page_content: str
            metadata: Dict[str, Any]


try:
    import pypdf
except ImportError:
    pypdf = None  # type: ignore[assignment]

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore[assignment]


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".json", ".csv", ".html", ".htm"}


class DocumentLoader:
    """Loads documents from a local directory (recursively) or a single file.

    Usage::

        loader = DocumentLoader()
        docs = loader.load_documents("path/to/docs")
        single = loader.load_single_file("path/to/report.pdf")
    """

    def __init__(self, recursive: bool = True) -> None:
        self.recursive = recursive

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def load_documents(
        self,
        path: str | Path,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """Load all supported files from *path* (directory) recursively.

        Args:
            path: directory to scan.
            extensions: whitelist of extensions (e.g. ``[".pdf", ".txt"]``).
                        Defaults to all ``SUPPORTED_EXTENSIONS``.

        Returns:
            List of ``Document`` objects with metadata.
        """
        base = Path(path).expanduser().resolve()
        if not base.exists():
            raise FileNotFoundError(f"Directory does not exist: {base}")
        if not base.is_dir():
            raise ValueError(f"Expected a directory, got a file: {base}")

        if extensions is None:
            exts = SUPPORTED_EXTENSIONS
        else:
            exts = {("." + e.lstrip(".")).lower() for e in extensions}

        documents: List[Document] = []
        for file_path in self._iter_files(base, exts):
            try:
                docs = self._load_file(file_path)
                documents.extend(docs)
            except Exception as exc:
                print(f"[DocumentLoader] Error reading {file_path}: {exc}")

        print(f"[DocumentLoader] Loaded {len(documents)} document(s) from {base}")
        return documents

    def load_single_file(self, path: str | Path) -> List[Document]:
        """Load a single file and return a list of Documents.

        For PDFs each page becomes a separate Document.  For everything
        else the list contains exactly one element.

        Raises ``ValueError`` if the file extension is not supported.
        """
        file_path = Path(path).expanduser().resolve()
        if not file_path.exists():
            raise FileNotFoundError(f"File does not exist: {file_path}")
        if not file_path.is_file():
            raise ValueError(f"Expected a file, got a directory: {file_path}")
        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )
        return self._load_file(file_path)

    # ------------------------------------------------------------------
    # Internal: file iteration
    # ------------------------------------------------------------------

    def _iter_files(self, base: Path, exts: set[str]) -> List[Path]:
        results: List[Path] = []
        if self.recursive:
            for root, _dirs, files in os.walk(base):
                for name in files:
                    p = Path(root) / name
                    if p.suffix.lower() in exts and not name.startswith("."):
                        results.append(p)
        else:
            for p in base.iterdir():
                if p.is_file() and p.suffix.lower() in exts and not p.name.startswith("."):
                    results.append(p)
        return sorted(results)

    # ------------------------------------------------------------------
    # Internal: load a single file into Document(s)
    # ------------------------------------------------------------------

    def _load_file(self, path: Path) -> List[Document]:
        ext = path.suffix.lower()

        if ext in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
            if not text.strip():
                return []
            return [self._make_doc(text, path, file_type=ext.lstrip("."))]

        if ext == ".pdf":
            return self._read_pdf(path)

        if ext == ".docx":
            return self._read_docx(path)

        if ext == ".json":
            return self._read_json(path)

        if ext == ".csv":
            return self._read_csv(path)

        if ext in {".html", ".htm"}:
            return self._read_html(path)

        raise ValueError(f"Unsupported format: {ext}")

    # ------------------------------------------------------------------
    # Format-specific readers
    # ------------------------------------------------------------------

    def _read_pdf(self, path: Path) -> List[Document]:
        if pypdf is None:
            raise ImportError("pypdf is required for PDF files: pip install pypdf")
        docs: List[Document] = []
        with path.open("rb") as fh:
            reader = pypdf.PdfReader(fh)
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                if not text.strip():
                    continue
                docs.append(
                    self._make_doc(text, path, file_type="pdf", page=page_num)
                )
        return docs

    def _read_docx(self, path: Path) -> List[Document]:
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for DOCX files: pip install python-docx"
            )
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        if not text.strip():
            return []
        return [self._make_doc(text, path, file_type="docx")]

    def _read_json(self, path: Path) -> List[Document]:
        """Read a JSON file.

        If the JSON is a list of objects, each object becomes a Document
        (serialized back to a JSON string).  Otherwise the whole file
        content is treated as a single Document.
        """
        raw = path.read_text(encoding="utf-8", errors="ignore")
        if not raw.strip():
            return []
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            # Treat as plain text if invalid JSON
            return [self._make_doc(raw, path, file_type="json")]

        if isinstance(data, list):
            docs: List[Document] = []
            for idx, item in enumerate(data):
                text = json.dumps(item, ensure_ascii=False, indent=2) if not isinstance(item, str) else item
                if text.strip():
                    docs.append(
                        self._make_doc(text, path, file_type="json", page=idx + 1)
                    )
            return docs if docs else [self._make_doc(raw, path, file_type="json")]

        # Single object / scalar
        return [self._make_doc(raw, path, file_type="json")]

    def _read_csv(self, path: Path) -> List[Document]:
        """Read a CSV file.

        Each row becomes a separate Document whose ``page_content`` is
        a key: value representation of that row (using the CSV header).
        If the CSV has no header or only one row, the entire file is
        returned as one Document.
        """
        raw = path.read_text(encoding="utf-8", errors="ignore")
        if not raw.strip():
            return []

        reader = csv.DictReader(io.StringIO(raw))
        fieldnames = reader.fieldnames
        if not fieldnames:
            return [self._make_doc(raw, path, file_type="csv")]

        docs: List[Document] = []
        for row_num, row in enumerate(reader, start=1):
            lines = [f"{k}: {v}" for k, v in row.items() if v]
            text = "\n".join(lines)
            if text.strip():
                docs.append(
                    self._make_doc(text, path, file_type="csv", page=row_num)
                )

        return docs if docs else [self._make_doc(raw, path, file_type="csv")]

    def _read_html(self, path: Path) -> List[Document]:
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self._parts: List[str] = []

            def handle_data(self, data: str) -> None:
                data = data.strip()
                if data:
                    self._parts.append(data)

            def get_text(self) -> str:
                return " ".join(self._parts)

        content = path.read_text(encoding="utf-8", errors="ignore")
        parser = TextExtractor()
        parser.feed(content)
        text = parser.get_text()
        if not text.strip():
            return []
        return [self._make_doc(text, path, file_type=path.suffix.lstrip(".").lower())]

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _make_doc(
        self,
        text: str,
        path: Path,
        *,
        file_type: str,
        page: Optional[int] = None,
    ) -> Document:
        stat = path.stat()
        meta: Dict[str, Any] = {
            "source": path.name,
            "file_path": str(path.resolve()),
            "file_type": file_type,
            "format": file_type,
            "mtime": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "size_bytes": stat.st_size,
            "content_hash": self._hash(text),
        }
        if page is not None:
            meta["page"] = page
        return Document(page_content=text, metadata=meta)

    @staticmethod
    def _hash(text: str) -> str:
        return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


class DocumentChangeTracker:
    """Save and compare document hashes between ingestion runs."""

    @staticmethod
    def save_state(docs: List[Document], path: str | Path) -> None:
        state: Dict[str, Any] = {}
        for doc in docs:
            meta = doc.metadata
            key = meta.get("file_path") or meta.get("source")
            if not key:
                continue
            state[key] = {
                "mtime": meta.get("mtime"),
                "size_bytes": meta.get("size_bytes"),
                "content_hash": meta.get("content_hash"),
            }

        out = Path(path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    @staticmethod
    def load_state(path: str | Path) -> Dict[str, Any]:
        p = Path(path)
        if not p.exists():
            return {}
        return json.loads(p.read_text(encoding="utf-8"))

    @staticmethod
    def diff(
        docs: List[Document],
        previous_state: Dict[str, Any],
    ) -> Dict[str, List[str]]:
        current: Dict[str, str] = {}
        for doc in docs:
            meta = doc.metadata
            key = meta.get("file_path") or meta.get("source")
            if key:
                current[key] = meta.get("content_hash", "")

        changes = {"new": [], "modified": [], "deleted": []}

        for key, content_hash in current.items():
            previous = previous_state.get(key)
            if previous is None:
                changes["new"].append(key)
            elif previous.get("content_hash") != content_hash:
                changes["modified"].append(key)

        for key in previous_state.keys():
            if key not in current:
                changes["deleted"].append(key)

        return changes
